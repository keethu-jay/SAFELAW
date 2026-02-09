#!/usr/bin/env python3
"""
Ingest Small Corpus (5 papers) into Supabase mini tables for specialized testing.

Reads docx files from backend/Small Corpus/:
  - *_paragraphs.docx: paragraph-level chunks
  - *_sentences.docx: sentence-level chunks

Inserts into:
  - corpus_documents_mini_paragraphs
  - corpus_documents_mini_sentences

Requires: Run supabase_mini_corpus_schema.sql in Supabase SQL Editor first.
Env: .env with SUPABASE_URL, SUPABASE_KEY, ISAACUS_API_KEY.
"""

import os
import re
import sys
from pathlib import Path

SCRIPT_DIR = Path(__file__).resolve().parent
BACKEND_DIR = SCRIPT_DIR.parent
SMALL_CORPUS_DIR = BACKEND_DIR / "Small Corpus"

_env_path = BACKEND_DIR / ".env"
if _env_path.exists():
    with open(_env_path) as f:
        for line in f:
            line = line.strip()
            if line and not line.startswith("#") and "=" in line:
                k, v = line.split("=", 1)
                os.environ[k.strip()] = v.strip()

sys.path.insert(0, str(BACKEND_DIR / "src"))
from embedding_helper import EmbeddingModel

# Map filename stem (without _paragraphs/_sentences) -> doc_id for display
DOC_ID_MAP = {
    "Donoghue_v_Stevenson": "Donoghue v Stevenson [1932] UKHL 100",
    "Kelly_v_Hennessy": "Ann Kelly v Fergus Hennessy [1995] 3 IR 253",
    "McGee_v_Attorney_General": "McGee v A.G. and Anor [1973] IESC 2",
    "McLoughlin_v_OBrian": "McLoughlin v O'Brian [1982] UKHL 3",
    "Norris_v_Attorney_General": "Norris v A.G. [1983] IESC 3",
}


def extract_paragraphs_from_docx(path: Path) -> list[str]:
    """Extract paragraphs from docx (one paragraph element = one chunk)."""
    try:
        from docx import Document
        doc = Document(path)
        paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        if not paras:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        t = cell.text.strip()
                        if t:
                            paras.append(t)
        return paras
    except ImportError:
        # Fallback: use zip + xml
        import zipfile
        import xml.etree.ElementTree as ET
        W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        paras = []
        try:
            with zipfile.ZipFile(path, "r") as z:
                with z.open("word/document.xml") as f:
                    root = ET.parse(f).getroot()
            for p_elem in root.iter(f"{{{W_NS}}}p"):
                texts = []
                for t_elem in p_elem.iter(f"{{{W_NS}}}t"):
                    if t_elem.text:
                        texts.append(t_elem.text)
                para = "".join(texts).strip()
                if para:
                    paras.append(para)
        except Exception as e:
            print(f"  Error reading {path.name}: {e}")
        return paras


def extract_sentences_from_docx(path: Path) -> list[str]:
    """Extract sentences from docx (split paragraphs by sentence boundaries)."""
    paras = extract_paragraphs_from_docx(path)
    sentences = []
    for p in paras:
        # Split on . ! ? followed by space or end
        parts = re.split(r'(?<=[.!?])\s+', p)
        for s in parts:
            s = s.strip()
            if s and len(s) >= 10:
                sentences.append(s)
    return sentences


def stem_to_doc_id(stem: str) -> str:
    """Get doc_id from filename stem (e.g. Donoghue_v_Stevenson)."""
    return DOC_ID_MAP.get(stem, stem.replace("_", " "))


def main():
    if not SMALL_CORPUS_DIR.exists():
        print(f"Small Corpus folder not found: {SMALL_CORPUS_DIR}")
        sys.exit(1)

    try:
        from supabase import create_client
    except ImportError:
        print("Install supabase: pip install supabase")
        sys.exit(1)

    supabase_url = os.environ.get("SUPABASE_URL")
    supabase_key = os.environ.get("SUPABASE_KEY")
    if not supabase_url or not supabase_key:
        print("Set SUPABASE_URL and SUPABASE_KEY in .env")
        sys.exit(1)

    model = EmbeddingModel(model_name="kanon-2-embedder")
    if not model.client:
        print("Set ISAACUS_API_KEY in .env")
        sys.exit(1)

    supabase = create_client(supabase_url, supabase_key)

    def embed_batch(texts: list[str], max_chars: int = 8000) -> list:
        if not texts:
            return []
        truncated = [t[:max_chars] + ("..." if len(t) > max_chars else "") for t in texts if t and str(t).strip()]
        if not truncated:
            import numpy as np
            return [np.zeros(1792) for _ in texts]
        try:
            resp = model.client.embeddings.create(
                model=model.model_name, texts=truncated, task="retrieval/document"
            )
            import numpy as np
            return [np.array(e.embedding) for e in resp.embeddings]
        except Exception as e:
            print(f"  Embed failed: {e}")
            return []

    BATCH_SIZE = 20

    # ---- Ingest paragraphs ----
    para_files = sorted(SMALL_CORPUS_DIR.glob("*_paragraphs.docx"))
    print(f"\nIngesting {len(para_files)} paragraph files into corpus_documents_mini_paragraphs...")

    # Clear existing (optional - comment out to append)
    try:
        supabase.table("corpus_documents_mini_paragraphs").delete().neq("id", 0).execute()
        print("  Cleared existing mini_paragraphs rows.")
    except Exception as e:
        print(f"  Note: Could not clear mini_paragraphs: {e}")

    global_idx = 0
    for path in para_files:
        stem = path.stem.replace("_paragraphs", "")
        doc_id = stem_to_doc_id(stem)
        chunks = extract_paragraphs_from_docx(path)
        chunks = [c for c in chunks if len(c) >= 10]
        if not chunks:
            print(f"  Skip {path.name}: no paragraphs")
            continue

        for i in range(0, len(chunks), BATCH_SIZE):
            batch = chunks[i : i + BATCH_SIZE]
            embs = embed_batch(batch)
            if len(embs) != len(batch):
                print(f"  Embed mismatch for {path.name} batch")
                continue
            rows = []
            for j, (text, emb) in enumerate(zip(batch, embs)):
                rows.append({
                    "doc_id": doc_id,
                    "text": text,
                    "section_title": "Main",
                    "section_number": str(j + 1),
                    "sentence_index": i + j,
                    "global_index": global_idx,
                    "court": "UKSC",
                    "decision": "majority",
                    "embedding": emb.tolist(),
                })
                global_idx += 1
            try:
                supabase.table("corpus_documents_mini_paragraphs").insert(rows).execute()
            except Exception as e:
                print(f"  Insert error: {e}")
                for r in rows:
                    try:
                        supabase.table("corpus_documents_mini_paragraphs").insert([r]).execute()
                    except Exception as e2:
                        print(f"    Failed: {e2}")
        print(f"  Ingested {path.name}: {len(chunks)} paragraphs")

    # ---- Ingest sentences ----
    sent_files = sorted(SMALL_CORPUS_DIR.glob("*_sentences.docx"))
    print(f"\nIngesting {len(sent_files)} sentence files into corpus_documents_mini_sentences...")

    try:
        supabase.table("corpus_documents_mini_sentences").delete().neq("id", 0).execute()
        print("  Cleared existing mini_sentences rows.")
    except Exception as e:
        print(f"  Note: Could not clear mini_sentences: {e}")

    global_idx = 0
    for path in sent_files:
        stem = path.stem.replace("_sentences", "")
        doc_id = stem_to_doc_id(stem)
        chunks = extract_sentences_from_docx(path)
        if not chunks:
            print(f"  Skip {path.name}: no sentences")
            continue

        for i in range(0, len(chunks), BATCH_SIZE):
            batch = chunks[i : i + BATCH_SIZE]
            embs = embed_batch(batch)
            if len(embs) != len(batch):
                print(f"  Embed mismatch for {path.name} batch")
                continue
            rows = []
            for j, (text, emb) in enumerate(zip(batch, embs)):
                rows.append({
                    "doc_id": doc_id,
                    "text": text,
                    "section_title": "Main",
                    "section_number": str(j + 1),
                    "sentence_index": i + j,
                    "global_index": global_idx,
                    "court": "UKSC",
                    "decision": "majority",
                    "embedding": emb.tolist(),
                })
                global_idx += 1
            try:
                supabase.table("corpus_documents_mini_sentences").insert(rows).execute()
            except Exception as e:
                print(f"  Insert error: {e}")
                for r in rows:
                    try:
                        supabase.table("corpus_documents_mini_sentences").insert([r]).execute()
                    except Exception as e2:
                        print(f"    Failed: {e2}")
        print(f"  Ingested {path.name}: {len(chunks)} sentences")

    print("\n✅ Small corpus ingestion complete.")


if __name__ == "__main__":
    main()
