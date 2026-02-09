#!/usr/bin/env python3
"""
Convert all docx and rtf files to a consistent XML format for semantic similarity.

Reads:
- Test Files/xml docs/*.docx (Donoghue, McLoughlin, Ann Kelly)
- Test Files/raw docs/*.rtf (2.rtf, 3.rtf)

Outputs:
- Test Files/xml docs/*.xml (one per source, same folder)

Uses a simple, readable XML structure compatible with run_baseline_corpus_studio.py
(extract_full_text uses root.itertext() so any well-formed XML with text works).

Requires: pip install python-docx striprtf (or uses built-in fallbacks if not installed)
"""

import html
import re
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

SCRIPT_DIR = Path(__file__).resolve().parent
XML_DOCS_DIR = SCRIPT_DIR / "xml docs"
RAW_DOCS_DIR = SCRIPT_DIR / "raw docs"

# Map source filename -> output XML filename (without path)
# RTF files 2.rtf and 3.rtf: derive title from content or use friendly name
OUTPUT_NAMES = {
    "Donoghue_v_Stevenson__1932__UKHL_100__26_May_1932_(pdfgear.com).docx": "Donoghue v Stevenson [1932] UKHL 100.xml",
    "mcloughlin-v-obrian(pdfgear.com).docx": "McLoughlin v O'Brian [1982] UKHL 3.xml",
    "Ann Kelly, Plaintiff, v Fergus Hennessy, Defendant [1995] 3 IR 253.docx": "Ann Kelly v Fergus Hennessy [1995] 3 IR 253.xml",
    "2.rtf": "McGee v A.G. and Anor [1973] IESC 2.xml",
    "3.rtf": "Norris v A.G. [1983] IESC 3.xml",
}


def escape_xml(text: str) -> str:
    return html.escape(text or "", quote=False)


def text_to_paragraphs(text: str) -> list[str]:
    """Split text into paragraphs (non-empty blocks separated by blank lines)."""
    if not text or not text.strip():
        return []
    paras = re.split(r"\n\s*\n", text.strip())
    return [p.strip() for p in paras if p.strip()]


def _docx_via_docx(path: Path) -> list[str]:
    """Extract paragraphs using python-docx."""
    from docx import Document
    doc = Document(path)
    paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    if not paras:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t:
                        paras.append(t)
    return paras


def _docx_via_zip(path: Path) -> list[str]:
    """Fallback: extract paragraphs from docx (zip) word/document.xml using standard lib."""
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    paras = []
    try:
        with zipfile.ZipFile(path, "r") as z:
            with z.open("word/document.xml") as f:
                root = ET.parse(f).getroot()
        for p_elem in root.iter(f"{{{W_NS}}}p"):
            texts = []
            for t_elem in p_elem.iter(f"{{{W_NS}}}t"):
                if t_elem.text:
                    texts.append(t_elem.text)
            para = "".join(texts).strip()
            if para:
                paras.append(para)
    except Exception:
        pass
    return paras


def docx_to_paragraphs(path: Path) -> list[str]:
    """Extract paragraph text from a docx file."""
    try:
        return _docx_via_docx(path)
    except ImportError:
        paras = _docx_via_zip(path)
        if not paras:
            print("  (pip install python-docx for better extraction)")
        return paras if paras else [""]
    except Exception as e:
        print(f"  Error reading docx: {e}")
        paras = _docx_via_zip(path)
        return paras if paras else [""]


def _rtf_strip_plain(raw: str) -> str:
    """Simple RTF-to-text using regex (fallback when striprtf not installed)."""
    text = raw
    # Replace \par with newline to preserve paragraph breaks
    text = re.sub(r"\\par\b", "\n", text, flags=re.IGNORECASE)
    # Remove RTF control words: \word or \word123 or \word-123
    text = re.sub(r"\\[a-zA-Z]+\-?\d*\s?", " ", text)
    text = re.sub(r"\\[{} \\]", " ", text)
    # Decode hex escapes \'xx (e.g. \'93 = ', \'94 = ")
    def hex_repl(m):
        try:
            return chr(int(m.group(1), 16))
        except ValueError:
            return ""
    text = re.sub(r"\\'([0-9a-fA-F]{2})", hex_repl, text)
    # Remove braces and backslashes; keep text
    text = text.replace("{", " ").replace("}", " ").replace("\\", " ")
    # Normalize whitespace
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n\s*\n+", "\n\n", text).strip()
    return text


def rtf_to_paragraphs(path: Path) -> list[str]:
    """Extract paragraph text from an RTF file."""
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        raw = f.read()
    try:
        from striprtf.striprtf import rtf_to_text
        text = rtf_to_text(raw)
    except ImportError:
        text = _rtf_strip_plain(raw)
        if not text.strip():
            print("  (pip install striprtf for better extraction)")
    except Exception as e:
        print(f"  RTF error: {e}")
        text = _rtf_strip_plain(raw)
    paras = text_to_paragraphs(text)
    return paras if paras else [""]


def rtf_title_from_content(path: Path) -> str | None:
    """Try to extract case title from RTF info block or first heading."""
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            raw = f.read()
        # RTF \title{...} or \info{\title ...}
        m = re.search(r"\\title\s*([^{}]+)|\\title\s*\{([^{}]+)\}", raw, re.IGNORECASE)
        if m:
            t = (m.group(1) or m.group(2) or "").strip()
            if t:
                return t
        # Fallback: first line of readable text
        from striprtf.striprtf import rtf_to_text
        text = rtf_to_text(raw)
        first = text.strip().split("\n")[0].strip()[:80]
        if first:
            return first
    except Exception:
        pass
    return None


def build_xml(title: str, paragraphs: list[str]) -> str:
    """Build a simple, readable XML document."""
    lines = [
        '<?xml version="1.0" encoding="UTF-8"?>',
        "<document>",
        "  <meta>",
        f"    <title>{escape_xml(title)}</title>",
        "  </meta>",
        "  <body>",
    ]
    for p in paragraphs:
        if p:
            lines.append(f"    <p>{escape_xml(p)}</p>")
    lines.append("  </body>")
    lines.append("</document>")
    return "\n".join(lines)


def main():
    XML_DOCS_DIR.mkdir(parents=True, exist_ok=True)

    # 1) Convert docx in xml docs
    for f in sorted(XML_DOCS_DIR.glob("*.docx")):
        out_name = OUTPUT_NAMES.get(f.name)
        if not out_name:
            out_name = f.stem + ".xml"
        out_path = XML_DOCS_DIR / out_name
        if out_path.exists():
            print(f"Skip (exists): {out_name}")
            continue
        print(f"Convert docx: {f.name} -> {out_name}")
        paras = docx_to_paragraphs(f)
        title = Path(out_name).stem
        xml = build_xml(title, paras)
        out_path.write_text(xml, encoding="utf-8")
        print(f"  Saved: {out_name} ({len(paras)} paragraphs)")

    # 2) Convert rtf in raw docs
    for f in sorted(RAW_DOCS_DIR.glob("*.rtf")):
        out_name = OUTPUT_NAMES.get(f.name)
        if not out_name:
            out_name = f.stem + ".xml"
        out_path = XML_DOCS_DIR / out_name
        if out_path.exists():
            print(f"Skip (exists): {out_name}")
            continue
        print(f"Convert rtf: {f.name} -> {out_name}")
        paras = rtf_to_paragraphs(f)
        title = rtf_title_from_content(f) or Path(out_name).stem
        xml = build_xml(title, paras)
        out_path.write_text(xml, encoding="utf-8")
        print(f"  Saved: {out_name} ({len(paras)} paragraphs)")

    print("\nDone. All five docs are now XML in Test Files/xml docs.")


if __name__ == "__main__":
    main()
